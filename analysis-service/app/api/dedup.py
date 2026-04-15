"""
文档查重 API 路由
"""

import io
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel, Field

from app.services.dedup import dedup_service

logger = logging.getLogger(__name__)

router = APIRouter()


class DedupResult(BaseModel):
    """查重结果项"""
    id: int
    fileName: str
    similarity: float
    type: str  # exact / semantic / core
    matchedDoc: str
    matchedContent: str
    uploadedContent: str


class DedupStats(BaseModel):
    """查重统计"""
    total_matches: int = 0
    exact_count: int = 0
    semantic_count: int = 0
    core_count: int = 0
    high_similarity: int = 0
    medium_similarity: int = 0
    low_similarity: int = 0


class DedupResponse(BaseModel):
    """查重响应"""
    results: list[DedupResult] = Field(default_factory=list)
    stats: DedupStats = Field(default_factory=DedupStats)


class ImportResponse(BaseModel):
    """导入响应"""
    total: int
    imported: int
    failed: int
    errors: list[str] = Field(default_factory=list)
    fingerprint_count: int = 0


@router.post("/check", response_model=DedupResponse)
async def check_documents(files: list[UploadFile] = File(...)):
    """
    上传文件进行查重

    支持 Word (.docx)、PDF (.pdf) 格式，可同时上传多个文件。
    系统将与文档库进行三级查重比对：完全一致、语义相似、核心内容相同。
    """
    if not files:
        raise HTTPException(status_code=400, detail="请上传至少一个文件")

    contents = []
    for file in files:
        if not file.filename:
            continue

        filename = file.filename.lower()
        raw_content = await file.read()

        try:
            text = ""

            if filename.endswith(".txt"):
                text = raw_content.decode("utf-8")

            elif filename.endswith(".docx"):
                try:
                    from docx import Document
                    doc = Document(io.BytesIO(raw_content))
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                except ImportError:
                    raise HTTPException(status_code=500, detail="Word 文件解析依赖未安装")

            elif filename.endswith(".pdf"):
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(raw_content))
                    text = "\n".join([page.extract_text() or "" for page in reader.pages])
                except ImportError:
                    raise HTTPException(status_code=500, detail="PDF 文件解析依赖未安装")

            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"不支持的文件格式: {file.filename}，请上传 .docx、.pdf 或 .txt 文件",
                )

            if text.strip():
                contents.append({
                    "filename": file.filename,
                    "content": text,
                })

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"文件解析失败 {file.filename}: {str(e)}")
            raise HTTPException(status_code=400, detail=f"文件解析失败: {file.filename}")

    if not contents:
        raise HTTPException(status_code=400, detail="未能从上传文件中提取到有效内容")

    try:
        result = await dedup_service.check_documents(contents)
        return result
    except Exception as e:
        logger.error(f"查重失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"查重服务异常: {str(e)}")


@router.get("/result/{task_id}")
async def get_result(task_id: str):
    """
    获取查重结果详情

    用于异步查重场景，通过任务 ID 获取查重结果。
    当前版本为同步查重，此接口预留。
    """
    # 预留接口，当前版本查重为同步操作
    raise HTTPException(
        status_code=404,
        detail="当前版本为同步查重，请直接使用 /check 接口获取结果",
    )


@router.post("/import", response_model=ImportResponse)
async def import_documents(files: list[UploadFile] = File(...)):
    """
    增量导入文档到文档库

    上传文档到查重文档库，用于后续查重比对。
    """
    if not files:
        raise HTTPException(status_code=400, detail="请上传至少一个文件")

    documents = []
    for file in files:
        if not file.filename:
            continue

        filename = file.filename.lower()
        raw_content = await file.read()

        try:
            text = ""

            if filename.endswith(".txt"):
                text = raw_content.decode("utf-8")
            elif filename.endswith(".docx"):
                from docx import Document
                doc = Document(io.BytesIO(raw_content))
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            elif filename.endswith(".pdf"):
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(raw_content))
                text = "\n".join([page.extract_text() or "" for page in reader.pages])

            if text.strip():
                documents.append({
                    "filename": file.filename,
                    "content": text,
                })

        except Exception as e:
            logger.warning(f"导入文件解析失败 {file.filename}: {str(e)}")

    if not documents:
        raise HTTPException(status_code=400, detail="未能从上传文件中提取到有效内容")

    try:
        result = await dedup_service.import_documents(documents)
        return result
    except Exception as e:
        logger.error(f"文档导入失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档导入失败: {str(e)}")
